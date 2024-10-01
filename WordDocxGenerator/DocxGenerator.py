import docx.text.paragraph
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import RGBColor, Pt


class DocxGenerator:
    def __init__(self):
        self.__doc_handle = Document()
        self.__generate_path = None

    def set_path(self, path: str):
        self.__generate_path = path

    def add_paragraph(self, text: str):
        self.__doc_handle.styles['Normal'].font.name = u'宋体'
        self.__doc_handle.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
        self.__doc_handle.styles['Normal'].font.size = Pt(10.5)
        self.__doc_handle.styles['Normal'].font.color.rgb = RGBColor(0, 0, 0)
        return self.__doc_handle.add_paragraph(text)

    @staticmethod
    def append_paragraph(para: docx.text.paragraph.Paragraph, text: str, rgb: list[int]):
        run = para.add_run(text)
        run.font.name = u'宋体'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
        run.font.size = Pt(10.5)
        run.font.color.rgb = RGBColor(rgb[0], rgb[1], rgb[2])

    def add_heading(self, text: str, level: int, req_center: bool):
        head = self.__doc_handle.add_heading("", level=level)
        if req_center:
            head.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = head.add_run(text)
        run.font.name = u'楷体'
        if level == 1:
            run.font.size = Pt(18)
        run._element.rPr.rFonts.set(qn('w:eastAsia'), u'楷体')
        run.font.color.rgb = RGBColor(0, 0, 0)

    def save_this(self):
        self.__doc_handle.save(self.__generate_path)
