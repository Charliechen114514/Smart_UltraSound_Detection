import docx.text.paragraph
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import RGBColor, Pt, Cm


class BasicInfoCreator:
    @staticmethod
    def infos_template() -> str:
        return \
            ("姓名:            性别：		        年龄：	        超声号：                \n"
             "科别：         	        门诊号：    	                        床位：     \n"
             "检查项目：胎儿彩色多普勒超声检查")

    @staticmethod
    def dates_and_docter() -> str:
        return "诊断医师：	报告日期："

class TableTextProvider:
    HIGH_SYM = "↑"
    LOW_SYM = "↓"
    def __init__(self, params: list[float]):
        self.__params = params
        self.__period = True

    def set_period(self, st: bool):
        self.__period = st

    @staticmethod
    def __check_status_param(value: float, lower: float, upper: float):
        # valid params
        if value < lower:
            return -1
        elif value > upper:
            return 1
        return 0

    @staticmethod
    def __gain_bonus_scripts(res: int, nor_range: str) -> str:
        if res == -1:
            return TableTextProvider.LOW_SYM + nor_range
        elif res == 1:
            return TableTextProvider.HIGH_SYM + nor_range
        return nor_range

    def fetch_text(self, index):
        # ED SD don't need bonus mark
        if index == 1 or index == 0:
            return str(self.__params[index]) + "cm/s"

        # S/D
        if index == 2:
            if self.__period:
                return str(self.__params[index]) + \
                    TableTextProvider.__gain_bonus_scripts(
                        TableTextProvider.__check_status_param(self.__params[index],
                                                               3, 4.5),
                        "（3.0-4.5）"
                    )
            else:
                return str(self.__params[index]) + \
                    TableTextProvider.__gain_bonus_scripts(
                        TableTextProvider.__check_status_param(self.__params[index],
                                                               1.7, 3),
                        "（1.7-3.0）"
                    )

        # PI
        if index == 3:
            if self.__period:
                return str(self.__params[index]) + \
                    TableTextProvider.__gain_bonus_scripts(
                        TableTextProvider.__check_status_param(self.__params[index],
                                                               0.86, 1.2),
                        "（0.86-1.2）"
                    )
            else:
                return str(self.__params[index]) + \
                    TableTextProvider.__gain_bonus_scripts(
                        TableTextProvider.__check_status_param(self.__params[index],
                                                               0.62, 1.24),
                        "（0.62-1.24）"
                    )

        # RI
        if index == 4:
            if self.__period:
                return str(self.__params[index]) + \
                    TableTextProvider.__gain_bonus_scripts(
                        TableTextProvider.__check_status_param(self.__params[index],
                                                               0.6, 0.7),
                        "（0.6 - 0.7）"
                    )
            else:
                return str(self.__params[index]) + \
                    TableTextProvider.__gain_bonus_scripts(
                        TableTextProvider.__check_status_param(self.__params[index],
                                                               0.45, 0.65),
                        "（0.45 - 0.65）"
                    )

        if index == 5:
            return str(self.__params[index]) + \
                TableTextProvider.__gain_bonus_scripts(
                    TableTextProvider.__check_status_param(
                        abs(self.__params[index]),30, 60),
                    "(30-60cm/s)"
                )

        if index == 6:
            if self.__period:
                return str(self.__params[index]) + \
                    TableTextProvider.__gain_bonus_scripts(
                        TableTextProvider.__check_status_param(
                            abs(self.__params[index]), 20, 60),
                        "(20-60cm/s)"
                    )
            else:
                return str(self.__params[index]) + \
                    TableTextProvider.__gain_bonus_scripts(
                        TableTextProvider.__check_status_param(
                            abs(self.__params[index]), 30, 70),
                        "(30-70cm/s)"
                    )

        if index == 7:
            return str(self.__params[index]) + \
                TableTextProvider.__gain_bonus_scripts(
                    TableTextProvider.__check_status_param(
                        abs(self.__params[index]), 120, 160),
                    "(120-160bpm)"
                )


class ReportDocxGeneratorCore:
    COL = 4
    ROW = 5
    TABLE_HEADERS = ["名称", "测值", "名称", "测值"]
    INDEX_TEXT = ["PS", "ED", "S/D", "PI", "RI", "MD", "TAMax", "HR"]
    LINES = "__________________________________________________________________________________"

    def __init__(self):
        self.__doc_handle = Document()
        self.__generate_path = None

    def set_path(self, path: str):
        self.__generate_path = path

    def add_paragraph(self, text: str):
        para = self.__doc_handle.add_paragraph()
        return self.append_paragraph(para, text)

    def append_paragraph(self, para: docx.text.paragraph.Paragraph, text: str):
        run = para.add_run(text)
        run.font.name = u'楷体'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), u'楷体')
        run.font.size = Pt(10.5)
        run.font.color.rgb = RGBColor(0, 0, 0)
        return run

    def draw_infos(self, para: docx.text.paragraph.Paragraph):
        run = para.add_run(ReportDocxGeneratorCore.LINES)
        run.font.name = u'楷体'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), u'楷体')
        run.font.size = Pt(10.5)
        run.font.color.rgb = RGBColor(0, 0, 0)
        para = self.__doc_handle.add_paragraph()
        run = self.append_paragraph(para, BasicInfoCreator.infos_template())
        run.add_text(ReportDocxGeneratorCore.LINES)
        return run

    def __set_cell_text(self, table, index_row: int, index_col: int, index_text: str, text_text: str):
        index_cell = table.cell(index_row, index_col)
        index_cell.text = index_text
        edit_cell = table.cell(index_row, index_col + 1)
        edit_cell.text = text_text
        for each_para in index_cell.paragraphs:
            run = each_para.runs[0]
            run.font.name = u'楷体'
            run.font.bold = True
            run._element.rPr.rFonts.set(qn('w:eastAsia'), u'楷体')
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(0, 0, 0)
        for each_para in edit_cell.paragraphs:
            run = each_para.runs[0]
            run.font.name = u'宋体'
            run.font.bold = True
            run._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(0, 0, 0)

    def draw_table(self, infos: list[float]):
        if len(infos) != 8:
            return
        table_text_provider = TableTextProvider(infos)
        table = self.__doc_handle.add_table(ReportDocxGeneratorCore.ROW, ReportDocxGeneratorCore.COL)
        table.style = "Table Grid"
        table.autofit = True
        # init titles
        for i in range(ReportDocxGeneratorCore.COL):
            each_cell = table.cell(0, i)
            each_cell.text = ReportDocxGeneratorCore.TABLE_HEADERS[i]
            for each_para in each_cell.paragraphs:
                run = each_para.runs[0]
                run.font.name = u'宋体'
                run.font.bold = True
                run._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
                run.font.size = Pt(10)
                run.font.color.rgb = RGBColor(0, 0, 0)
        row_display = 0
        # init datas
        for i in range(8):
            if i <= 3:
                self.__set_cell_text(table, (row_display + 1) % 5,
                                     0, ReportDocxGeneratorCore.INDEX_TEXT[i], table_text_provider.fetch_text(i))
                row_display += 1
            else:
                if row_display == 4:
                    row_display = 0
                self.__set_cell_text(table, (row_display + 1) % 5,
                                     2, ReportDocxGeneratorCore.INDEX_TEXT[i],
                                        table_text_provider.fetch_text(i))
                row_display += 1

    def gain_blank_para(self):
        return self.__doc_handle.add_paragraph()

    def add_heading(self, text: str, level: int, req_center: bool, para: docx.text.paragraph.Paragraph):
        run = para.add_run(text)
        if req_center:
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run.font.name = u'楷体'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), u'楷体')
        size = Pt(10.5)
        if level == 1:
            size = Pt(16)
        elif level == 3:
            size = Pt(12)
        run.font.size = size
        run.font.color.rgb = RGBColor(0, 0, 0)

    def insert_pic(self, image_path: str, run, para: docx.text.paragraph.Paragraph):
        # 图片居中设置
        para = self.__doc_handle.add_paragraph()
        run = para.add_run()
        run.add_picture(image_path, width=Cm(8.98), height=Cm(6.73))
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        # 使用较小的段落间距以减小图片和表格之间的空白
        paragraph_format = para.paragraph_format
        paragraph_format.space_after = 0  # 设置段落后间距为0，避免图片与表格之间有空白
        return para

    def draw_fin(self, para: docx.text.paragraph.Paragraph):
        section = self.__doc_handle.sections[0]

        # 获取该section的页脚对象
        footer = section.footer

        # 向页脚添加内容（在页脚的第一个段落中插入文字）
        footer_paragraph = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        run = footer_paragraph.add_run("________________________________________________________________________"
                                       "__________\n" + BasicInfoCreator.dates_and_docter())
        run.font.name = u'楷体'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), u'楷体')
        run.font.size = Pt(10.5)
        run.font.color.rgb = RGBColor(0, 0, 0)

    def save_this(self):
        self.__doc_handle.save(self.__generate_path)
